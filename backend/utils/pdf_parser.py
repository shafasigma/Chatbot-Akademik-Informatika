import fitz  # PyMuPDF
import os
import re

try:
    from docx import Document  # python-docx untuk .docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx tidak terinstall. Install dengan: pip install python-docx")

try:
    import zipfile  # untuk ekstrak .doc format lama
    HAS_ZIP = True
except ImportError:
    HAS_ZIP = False

# ==============================
# KONFIGURASI FOLDER
# ==============================
PDF_DIR = os.path.join("backend", "uploads", "pdfs")
WORD_DIR = os.path.join("backend", "uploads", "words")  # Folder untuk file Word
TEXT_DIR = os.path.join("backend", "uploads", "texts")
os.makedirs(TEXT_DIR, exist_ok=True)
os.makedirs(WORD_DIR, exist_ok=True)


def split_sentences(text: str) -> list:
    """
    Pecah teks menjadi kalimat-kalimat individual.
    Menggunakan regex untuk mendeteksi titik, tanda tanya, tanda seru.
    """
    # Pisahkan berdasarkan . ! ? diikuti spasi
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    # Filter kalimat kosong
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


def extract_text_from_word(file_path: str) -> str:
    """
    Ekstrak teks dari file Word (.docx atau .doc).
    """
    if not HAS_DOCX:
        raise ImportError("python-docx tidak terinstall. Install dengan: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File Word tidak ditemukan: {file_path}")
    
    print(f"📘 Membaca file Word: {file_path}")
    
    try:
        doc = Document(file_path)
        full_text = ""
        
        # Ekstrak teks dari semua paragraf
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"
        
        # Ekstrak teks dari tabel jika ada
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + " "
                full_text += "\n"
        
        return full_text
    
    except Exception as e:
        raise Exception(f"❌ Gagal membaca file Word: {str(e)}")


def save_word_to_txt(word_path: str) -> str:
    """
    Ekstrak teks dari file Word ke file .txt dan pecah per kalimat.
    """
    if not os.path.exists(word_path):
        raise FileNotFoundError(f"❌ File Word tidak ditemukan: {word_path}")
    
    full_text = extract_text_from_word(word_path)
    
    # Bersihkan teks
    full_text = re.sub(r'\n+', ' ', full_text)  # Ganti newline dengan spasi
    full_text = re.sub(r'\s+', ' ', full_text)  # Hapus spasi ganda
    
    # Pecah per kalimat
    sentences_list = split_sentences(full_text)
    
    # Simpan hasil ke file teks
    txt_filename = os.path.splitext(os.path.basename(word_path))[0] + ".txt"
    txt_path = os.path.join(TEXT_DIR, txt_filename)
    with open(txt_path, "w", encoding="utf-8") as f:
        for idx, sentence in enumerate(sentences_list, 1):
            f.write(f"{idx}. {sentence}\n")
    
    print(f"✅ File Word berhasil diubah ke teks dan disimpan: {txt_path}")
    print(f"📄 Total kalimat ditemukan: {len(sentences_list)}")
    return txt_path



def save_pdf_to_txt(pdf_path: str) -> str:
    """
    Ekstrak teks dari PDF ke file .txt dan pecah per kalimat.
    Digunakan oleh admin_routes saat upload PDF baru.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"❌ File PDF tidak ditemukan: {pdf_path}")

    print(f"📘 Membaca PDF: {pdf_path}")
    doc = fitz.open(pdf_path)
    sentences_list = []

    # Gabung teks dari semua halaman
    full_text = ""
    for i, page in enumerate(doc):
        full_text += "\n" + page.get_text("text")
    doc.close()

    # Bersihkan teks
    full_text = re.sub(r'\n+', ' ', full_text)  # Ganti newline dengan spasi
    full_text = re.sub(r'\s+', ' ', full_text)  # Hapus spasi ganda

    # Pecah per kalimat
    sentences_list = split_sentences(full_text)

    # Simpan hasil ke file teks
    txt_filename = os.path.splitext(os.path.basename(pdf_path))[0] + ".txt"
    txt_path = os.path.join(TEXT_DIR, txt_filename)
    with open(txt_path, "w", encoding="utf-8") as f:
        for idx, sentence in enumerate(sentences_list, 1):
            f.write(f"{idx}. {sentence}\n")

    print(f"✅ PDF berhasil diubah ke teks dan disimpan: {txt_path}")
    print(f"📄 Total kalimat ditemukan: {len(sentences_list)}")
    return txt_path


if __name__ == "__main__":
    # 🔹 Contoh manual testing
    filename = "Peraturan-Rektor-No-25-Tahun-2020-Tentang-Penyelenggaraan-Kegiatan-Akademik-SEARCHABLE.pdf"
    full_path = os.path.join(PDF_DIR, filename)
    save_pdf_to_txt(full_path)
